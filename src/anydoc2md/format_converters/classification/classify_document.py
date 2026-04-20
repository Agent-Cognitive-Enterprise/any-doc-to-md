"""
DocumentTraits classifier — fast heuristic scan of a source document.

Returns a DocumentTraits frozen dataclass describing the document's
structural characteristics so the tournament can:
  - Weight adapter selection (e.g. prefer docling for image-heavy docs)
  - Prime the LLM judge ("look for OCR errors", "check table fidelity")
  - Weight QA checks (check_image_count_match matters more for image-heavy docs)

All detection is heuristic and best-effort; when a library is unavailable
or parsing fails, safe defaults (False / 0) are returned.

Supported formats: .pdf (PyMuPDF), .docx (python-docx),
                   .html/.htm (stdlib), .txt/.text (stdlib)
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

# ---------------------------------------------------------------------------
# Thresholds (exposed as module constants so callers can override)
# ---------------------------------------------------------------------------

# Scanned: words per page below this AND at least one image → likely scanned
SCANNED_WORDS_PER_PAGE_MAX: int = 60
# Image-heavy: unique images per page above this
IMAGE_HEAVY_IMAGES_PER_PAGE_MIN: float = 1.0
# Table-heavy: total tables across all pages above this
TABLE_HEAVY_MIN: int = 2
# Multi-column: pages with two or more distinct x-column zones ÷ total pages
MULTI_COLUMN_PAGE_FRACTION_MIN: float = 0.25
# Column gap: fraction of page width that must separate left and right text clusters
MULTI_COLUMN_GAP_FRACTION: float = 0.15
# Math: these patterns in extracted text suggest math notation
_MATH_PATTERNS = re.compile(
    r"\\frac|\\sum|\\int|\\alpha|\\beta|\\theta|\$[^$]{1,60}\$"
    r"|(?<!\w)(sin|cos|tan|log|lim)\s*[({]",
    re.IGNORECASE,
)
_HTML_IMG_RE = re.compile(r"<img\b", re.IGNORECASE)
_HTML_TABLE_RE = re.compile(r"<table\b", re.IGNORECASE)


# ---------------------------------------------------------------------------
# Result dataclass
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class DocumentTraits:
    """Structural characteristics of a source document."""

    file_type: str          # "pdf" | "docx" | "html" | "txt" | "unknown"
    page_count: int         # 0 for non-paginated formats
    image_count: int        # estimated unique image count
    table_count: int        # estimated table count
    word_count: int         # approximate total word count
    is_scanned: bool        # negligible text + images present → likely OCR needed
    is_image_heavy: bool    # many images relative to pages
    is_table_heavy: bool    # many structured tables
    is_multi_column: bool   # two-or-more column layout
    is_text_only: bool      # no images found
    has_math: bool          # math notation detected in extracted text

    def to_dict(self) -> dict:
        return {
            "file_type": self.file_type,
            "page_count": self.page_count,
            "image_count": self.image_count,
            "table_count": self.table_count,
            "word_count": self.word_count,
            "is_scanned": self.is_scanned,
            "is_image_heavy": self.is_image_heavy,
            "is_table_heavy": self.is_table_heavy,
            "is_multi_column": self.is_multi_column,
            "is_text_only": self.is_text_only,
            "has_math": self.has_math,
        }


def _unknown_traits(file_type: str = "unknown") -> DocumentTraits:
    return DocumentTraits(
        file_type=file_type, page_count=0, image_count=0, table_count=0,
        word_count=0, is_scanned=False, is_image_heavy=False,
        is_table_heavy=False, is_multi_column=False,
        is_text_only=True, has_math=False,
    )


# ---------------------------------------------------------------------------
# PDF classifier
# ---------------------------------------------------------------------------

def _is_multi_column_page(page, page_width: float) -> bool:
    """
    Heuristic: a page is multi-column when text blocks cluster into two
    distinct left and right zones separated by a gap > MULTI_COLUMN_GAP_FRACTION
    of the page width.
    """
    blocks = page.get_text("blocks")
    x0_values = [
        b[0] for b in blocks
        if b[6] == 0 and len(b[4].strip()) > 20  # text blocks with real content
    ]
    if len(x0_values) < 4:
        return False

    gap = MULTI_COLUMN_GAP_FRACTION * page_width
    left = [x for x in x0_values if x < page_width * 0.45]
    right = [x for x in x0_values if x > page_width * 0.45 + gap]

    return len(left) >= 2 and len(right) >= 2


def _classify_pdf(path: Path) -> DocumentTraits:
    try:
        import fitz
    except ImportError:
        return _unknown_traits("pdf")

    try:
        doc = fitz.open(str(path))
    except Exception:
        return _unknown_traits("pdf")

    import hashlib

    page_count = len(doc)
    seen_hashes: set[str] = set()
    total_words = 0
    table_count = 0
    full_text_parts: list[str] = []
    multi_col_pages = 0

    for page in doc:
        # Images (deduplicated by hash)
        for ref in page.get_images(full=True):
            try:
                base = doc.extract_image(ref[0])
                h = hashlib.sha256(base.get("image", b"")).hexdigest()[:16]
                seen_hashes.add(h)
            except Exception:
                pass

        # Text
        text = page.get_text("text", sort=True)
        total_words += len(text.split())
        full_text_parts.append(text)

        # Tables
        try:
            tabs = page.find_tables()
            table_count += len(tabs.tables)
        except Exception:
            pass

        # Multi-column
        if _is_multi_column_page(page, page.rect.width):
            multi_col_pages += 1

    doc.close()

    full_text = "\n".join(full_text_parts)
    image_count = len(seen_hashes)
    words_per_page = total_words / max(1, page_count)
    images_per_page = image_count / max(1, page_count)

    return DocumentTraits(
        file_type="pdf",
        page_count=page_count,
        image_count=image_count,
        table_count=table_count,
        word_count=total_words,
        is_scanned=(words_per_page < SCANNED_WORDS_PER_PAGE_MAX and image_count > 0),
        is_image_heavy=(images_per_page >= IMAGE_HEAVY_IMAGES_PER_PAGE_MIN),
        is_table_heavy=(table_count >= TABLE_HEAVY_MIN),
        is_multi_column=(multi_col_pages / max(1, page_count) >= MULTI_COLUMN_PAGE_FRACTION_MIN),
        is_text_only=(image_count == 0),
        has_math=bool(_MATH_PATTERNS.search(full_text)),
    )


# ---------------------------------------------------------------------------
# DOCX classifier
# ---------------------------------------------------------------------------

def _classify_docx(path: Path) -> DocumentTraits:
    try:
        import docx as _docx
    except ImportError:
        return _unknown_traits("docx")

    try:
        doc = _docx.Document(str(path))
    except Exception:
        return _unknown_traits("docx")

    # Count images via inline shapes
    image_count = sum(
        1 for p in doc.paragraphs
        for r in p.runs
        if r._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic"
        )
    )
    # Also check document body for drawing elements
    import re as _re
    body_xml = doc.element.body.xml
    # Rough counts via namespace tags
    image_count += len(_re.findall(r'<a:blip\b', body_xml))
    # Avoid double-counting: cap at unique blip refs
    image_count = len(set(_re.findall(r'r:embed="([^"]+)"', body_xml)))

    table_count = len(doc.tables)
    words = " ".join(
        p.text for p in doc.paragraphs if p.text.strip()
    )
    word_count = len(words.split())

    return DocumentTraits(
        file_type="docx",
        page_count=0,   # page count not available without rendering
        image_count=image_count,
        table_count=table_count,
        word_count=word_count,
        is_scanned=False,           # DOCX is never scanned
        is_image_heavy=(image_count >= 3),
        is_table_heavy=(table_count >= TABLE_HEAVY_MIN),
        is_multi_column=False,      # not detectable without rendering
        is_text_only=(image_count == 0),
        has_math=bool(_MATH_PATTERNS.search(words)),
    )


# ---------------------------------------------------------------------------
# HTML classifier
# ---------------------------------------------------------------------------

def _classify_html(path: Path) -> DocumentTraits:
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return _unknown_traits("html")

    image_count = len(_HTML_IMG_RE.findall(content))
    table_count = len(_HTML_TABLE_RE.findall(content))
    # Strip tags for word count
    text = re.sub(r"<[^>]+>", " ", content)
    word_count = len(text.split())

    return DocumentTraits(
        file_type="html",
        page_count=0,
        image_count=image_count,
        table_count=table_count,
        word_count=word_count,
        is_scanned=False,
        is_image_heavy=(image_count >= 5),
        is_table_heavy=(table_count >= TABLE_HEAVY_MIN),
        is_multi_column=False,  # CSS columns not detectable without rendering
        is_text_only=(image_count == 0),
        has_math=bool(_MATH_PATTERNS.search(text)),
    )


# ---------------------------------------------------------------------------
# Plain text classifier
# ---------------------------------------------------------------------------

def _classify_txt(path: Path) -> DocumentTraits:
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return _unknown_traits("txt")

    word_count = len(content.split())
    # Markdown-style table rows: lines with multiple | characters
    table_rows = sum(1 for line in content.splitlines() if line.count("|") >= 2)
    table_count = 1 if table_rows >= 3 else 0

    return DocumentTraits(
        file_type="txt",
        page_count=0,
        image_count=0,
        table_count=table_count,
        word_count=word_count,
        is_scanned=False,
        is_image_heavy=False,
        is_table_heavy=(table_count >= TABLE_HEAVY_MIN),
        is_multi_column=False,
        is_text_only=True,
        has_math=bool(_MATH_PATTERNS.search(content)),
    )


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

_CLASSIFIERS = {
    ".pdf":   _classify_pdf,
    ".docx":  _classify_docx,
    ".doc":   _classify_docx,
    ".html":  _classify_html,
    ".htm":   _classify_html,
    ".txt":   _classify_txt,
    ".text":  _classify_txt,
}


def classify(source_path: Path) -> DocumentTraits:
    """
    Classify a source document and return its structural traits.

    Falls back to _unknown_traits when the file type is unsupported or
    parsing fails.  Never raises.
    """
    suffix = source_path.suffix.lower()
    fn = _CLASSIFIERS.get(suffix)
    if fn is None:
        return _unknown_traits(suffix.lstrip(".") or "unknown")
    try:
        return fn(source_path)
    except Exception:
        return _unknown_traits(suffix.lstrip("."))
